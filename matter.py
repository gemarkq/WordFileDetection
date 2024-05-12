import utils
from utils import WordStatus
import docx
from nltk.tokenize import word_tokenize, sent_tokenize
import os
import sys
from zhon.hanzi import punctuation

class Matter:
    def __init__(self, path) -> None:
        self.path = path
        self.document = docx.Document(self.path)
        self.tables = self.document.tables
        self.paragraphs = []
        self.ExtractParagraphs()
    
    def ExtractParagraphs(self):
        for table in self.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self.paragraphs.append(para)

    def CheckClientName(self, client_name):
        first_name, last_name = client_name.split()
        wrong_flag = True
        for idx, para in enumerate(self.paragraphs):
            if not self.CheckName(para, first_name, last_name):
                print(f"[ERROR] paragraph: {idx}, {para.text}")
                wrong_flag = False
        return wrong_flag

    def CheckName(self, para, client_first_name, client_last_name):
        # Attorney ShuKui GAO
        # Attorney ShuKui GAO’s
        # Attoreny GAO
        # Attorney GAO’s
        wrong_first_name= []
        wrong_last_name = []
        wrong_attorney = []
        potentialc_client_name = [client_first_name + ' ' + client_last_name.upper(),
                                  client_first_name + ' ' + client_last_name.upper() + "'s",
                                  client_last_name.upper(),
                                  client_last_name.upper() + "'s"]
        words = word_tokenize(para.text)
        for i in range(len(words) - 1):
            # attorney check
            attorney_flag = utils.IsWordValid(words[i].lower(), "attorney")
            if attorney_flag == WordStatus.WORD_SPELL_CORRECT:
                pass
            elif attorney_flag == WordStatus.WORD_SPELL_WRONG:
                print(f"[ERROR] Attorney spell wrong")
                wrong_attorney.append(words[i])
            # first name check
            first_name_flag = utils.IsWordValid(words[i], client_first_name)
            if first_name_flag == WordStatus.WORD_SPELL_CORRECT:
                # last name check
                last_name_flag = utils.IsWordValid(words[i + 1], client_last_name)
                if last_name_flag == WordStatus.WORD_SPELL_CORRECT:
                    pass
                elif last_name_flag == WordStatus.WORD_SPELL_WRONG:
                    print(f"[ERROR] Client last name spell wrong, {words[i + 1]}")
                    wrong_last_name.append(client_last_name)
            elif first_name_flag == WordStatus.WORD_SPELL_WRONG:
                print(f"[ERROR] Client first name spell wrong, {words[i]}")
                wrong_first_name.append(words[i])

        first_length = len(wrong_first_name)
        last_length = len(wrong_last_name)
        att_length = len(wrong_attorney)

        if first_length != 0 or last_length != 0 or att_length != 0:
            return False
        return True

    def CheckPunctuation(self):
        pos = {}
        puncts = {}
        for idx, para in enumerate(self.paragraphs):
            sentences = sent_tokenize(para.text)
            for sentence in sentences:
                words = word_tokenize(sentence)
                # print(words)
                # print(sentence)
                for j, word in enumerate(words):
                    if not utils.IsValidWordPuctExclude(word):
                        # print(sentence)
                        pos.setdefault(idx, []).append(sentence)
                        if j > 0:
                            error_words = words[j-1] + words[j]
                        else:
                            error_words = words[j]
                        if j < len(words) - 1:
                            error_words += words[j+1]
                        puncts.setdefault(idx, []).append(error_words)
        
        if len(pos) != 0:
            for key in pos.keys():
                for idx, sentence in enumerate(pos[key]):
                    print(f"[ERROR] Use Chinese punctuation!!!, paragraph id: {key}, Error Words: {puncts[key][idx]}, Sentence: {sentence}")
            return False
        return True
    
    def FindAllChinese(self):
        pos = {}
        for idx, para in enumerate(self.paragraphs):
            sentences = sent_tokenize(para.text)
            for sentence in sentences:
                if utils.IsConatinChinese(sentence):
                    pos.setdefault(idx, []).append(sentence)
        
        if len(pos) != 0:
            for key in pos.keys():
                for sentence in pos[key]:
                    print(f"[Warning] Sentence: {sentence} in Paragraph {key} has Chinese, Please Check its Validity")
            return False
        return True


if __name__ == "__main__":
    document_path = input(r"Please Input document file path: ").strip()
    try:
        # document_path_ = document_path
        document_path = eval(document_path)
        print(f"Path Input: {document_path}\n")
    except Exception as e:
        print(f"Path Input: {document_path}\n")

    if not os.path.exists(document_path):
        print(f"{document_path} File Path does not exits. Please check!!!\n")
        input("Press enter to exit")
        sys.exit(0)
    
    client_name = input("Please Input Client Name (e.g. Shukui GAO): ")
    client_name_check = input("Please ReInput Client Name: ")
    client_name = client_name.strip()
    client_name_check = client_name_check.strip()
    while client_name != client_name_check:
        print(f"Client Name input mismatch. Please re-input the client name\n")
        client_name = input("Please Input Client Name (e.g. Shukui GAO): ")
        client_name_check = input("Please ReInput Client Name: ")
        client_name = client_name.strip()
        client_name_check = client_name_check.strip()
    
    utils.SeperateLine()
    print("CLIENT NAME CHEKC BEGIN")
    matter = Matter(document_path)
    if matter.CheckClientName(client_name):
        print(f"[Success] Client Name spell Correct!!! Congrats!")

    utils.SeperateLine()

    print("CHINESE PUNCTUATION CHECK BEGIN")
    if matter.CheckPunctuation():
        print(f"[Success] No punctuation error!!! Congrats!")
    
    utils.SeperateLine()

    print("CHINESE SENTENCE CHECK BEGIN")
    if matter.FindAllChinese():
        print(f"[Success] No Chinese Sentence Found!!! Congrats!")

    utils.SeperateLine()
    
    input("\n[Press enter to exit]")
    sys.exit(0)

